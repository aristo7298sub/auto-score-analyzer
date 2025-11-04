import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List
from app.models.score import StudentScore
from app.services.storage_service import StorageService

class ExportService:
    def __init__(self):
        self.storage_service = StorageService()

    async def export_to_excel(self, scores: List[StudentScore], file_path: str, original_filename: str = "") -> str:
        """
        导出成绩数据到Excel文件
        
        Args:
            scores: 学生成绩列表
            file_path: 导出文件路径
            original_filename: 原始上传的文件名（可选）
        """
        if not scores:
            raise ValueError("没有可导出的成绩数据")

        # 准备数据 - 简洁格式
        data = []
        for score in scores:
            row = {
                "学生姓名": score.student_name,
                "总分": score.total_score,
                "成绩分析": score.analysis or "暂无分析",
            }
            data.append(row)

        # 创建DataFrame并导出
        df = pd.DataFrame(data)
        
        # 使用openpyxl导出，设置列宽
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='成绩分析报告')
            
            # 获取工作表
            worksheet = writer.sheets['成绩分析报告']
            
            # 设置列宽
            worksheet.column_dimensions['A'].width = 15  # 学生姓名
            worksheet.column_dimensions['B'].width = 10  # 总分
            worksheet.column_dimensions['C'].width = 100  # 成绩分析
            
            # 设置成绩分析列自动换行
            from openpyxl.styles import Alignment
            for row in worksheet.iter_rows(min_row=2, max_row=worksheet.max_row, min_col=3, max_col=3):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical='top')
        
        return file_path

    async def export_to_word(self, scores: List[StudentScore], file_path: str, original_filename: str = "") -> str:
        """
        导出成绩数据到Word文件
        
        Args:
            scores: 学生成绩列表
            file_path: 导出文件路径
            original_filename: 原始上传的文件名（可选）
        """
        if not scores:
            raise ValueError("没有可导出的成绩数据")

        doc = Document()
        
        # 添加标题
        title = doc.add_heading('学生成绩分析报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加每个学生的成绩信息
        for idx, score in enumerate(scores):
            # 学生姓名
            student_p = doc.add_paragraph()
            student_run = student_p.add_run(f'学生：{score.student_name}')
            student_run.bold = True
            student_run.font.size = Pt(14)
            
            # 总分
            total_p = doc.add_paragraph(f'总分：{score.total_score}分')
            total_p.style = 'Normal'

            # 成绩分析（直接输出AI生成的完整分析文本）
            if score.analysis:
                analysis_p = doc.add_paragraph()
                analysis_run = analysis_p.add_run('成绩分析：')
                analysis_run.bold = True
                analysis_p.add_run(score.analysis)
                analysis_p.style = 'Normal'
            
            # 每个学生之间添加分隔
            if idx < len(scores) - 1:
                doc.add_paragraph()
                doc.add_paragraph('_' * 80)
                doc.add_paragraph()

        # 保存文档
        doc.save(file_path)
        return file_path 